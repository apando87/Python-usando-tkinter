import tkinter as tk
from tkinter import ttk, messagebox, simpledialog
from supabase import create_client, Client
from docx import Document
from tkinter import filedialog
import webbrowser
from edicion import EditaRegistro
from busqueda import BusquedaAvanzada
from PIL import Image, ImageTk
from tkinter.scrolledtext import ScrolledText  # <-- para la barra SOLO en Jurisprudencia
import math

import sys, os

try:
    from dotenv import load_dotenv
    load_dotenv()
except Exception:
    pass

SUPABASE_URL = os.environ.get("SUPABASE_URL")
SUPABASE_KEY = os.environ.get("SUPABASE_ANON_KEY")
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "")

if not SUPABASE_URL or not SUPABASE_KEY:
    try:
        messagebox.showerror(
            "Config faltante",
            "Definí SUPABASE_URL y SUPABASE_ANON_KEY en .env o variables de entorno."
        )
    except Exception:
        print("Faltan SUPABASE_URL y/o SUPABASE_ANON_KEY en el entorno.")
    raise SystemExit(1)

supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
# ─────────────────────────────────────────────────────────────────────────────

def resource_path(rel):
    """Permite encontrar archivos en modo .exe (PyInstaller) y en modo normal."""
    base = getattr(sys, '_MEIPASS', os.path.abspath("."))
    return os.path.join(base, rel)

BG_COLOR = "#FFFFFF"
FRAME_COLOR = "#F4F7FA"
BORDER_COLOR = "#E8EDF7"
TITULO_TEMA = ("Inter", 17, "bold")
TITULO_SUBTEMA = ("Inter", 13, "italic")
LABEL_FONT = ("Inter", 12, "bold")
INFO_FONT = ("Inter", 12)
BTN_ICONO = ("Segoe UI Symbol", 16)
BTN_COLOR = "#FFF"
BTN_TEXT_COLOR = "#1746A2"
BTN_HOVER_BG = "#1746A2"
BTN_HOVER_TEXT = "#FFF"
BTN_BORDER_COLOR = "#1746A2"
ICON_COPIAR = "\u29C9"   # ⧉
ICON_DESCARGAR = "\u2B07" # ⬇️
ICON_EDITAR = "\u270E"   # ✎
ICON_BORRAR = "\U0001F5D1"  # 🗑 (queda definido aunque ahora usamos imagen)

# Altura fija para que el área de Jurisprudencia quede alineada en todas las tarjetas
JURIS_ALTURA_LINEAS = 18  # subí o bajá este número si querés más/menos alto

# --- NUEVO: función para recolorear PNG a un color ---
def recolorear_icono(ruta_icono, color_hex, size=None):
    """
    Convierte un PNG (negro o monocromo) al color indicado (hex) manteniendo alfa.
    Opcionalmente cambia el tamaño (width, height).
    """
    img = Image.open(ruta_icono).convert("RGBA")
    if size:
        img = img.resize(size, Image.LANCZOS)

    datas = img.getdata()
    r = int(color_hex[1:3], 16)
    g = int(color_hex[3:5], 16)
    b = int(color_hex[5:7], 16)

    nueva = []
    for (pr, pg, pb, pa) in datas:
        # Si hay alfa (pixel visible), pintamos al color elegido preservando alfa
        if pa > 0:
            nueva.append((r, g, b, pa))
        else:
            nueva.append((pr, pg, pb, pa))
    img.putdata(nueva)
    return ImageTk.PhotoImage(img)

class VisualizadorBase(tk.Toplevel):
    def __init__(self, master=None):
        super().__init__(master)
        self.title("Visualización Compendio")
        ancho = 1450
        alto = 700
        x = (self.winfo_screenwidth() // 2) - (ancho // 2)
        y = (self.winfo_screenheight() // 2) - (alto // 2)
        self.geometry(f"{ancho}x{alto}+{x}+{y}")
        self.configure(bg=BG_COLOR)
        self.resizable(True, True)

        # --- DEBUG: bandera y toggle overlay ---
        self._debug_overlay_on = True
        self.bind("<F2>", lambda e: self._toggle_debug_overlay())

        # --- Cargar y guardar el icono de tacho recoloreado (para que no lo borre el GC) ---
        self.icono_tacho_azul = recolorear_icono(
            resource_path("tacho.png"),
            "#1746A2",
            size=(24, 24)
        )

        # PAGINACIÓN
        self.current_page = 1

        # --- Menú hamburguesa ---
        self.menu_popup = tk.Menu(
            self,
            tearoff=0,
            bg="#FFF",
            fg="#000",
            activebackground="#1746A2",
            activeforeground="#FFF"
        )
        self.menu_popup.add_command(label="Menú principal", command=self.volver_menu)
        self.menu_popup.add_command(label="Búsqueda avanzada", command=self.busqueda_avanzada)

        # --- Header ---
        header = tk.Frame(self, bg=BG_COLOR)
        header.pack(fill="x", pady=(28, 6))

        espaciador_izq = tk.Label(header, bg=BG_COLOR, width=4)
        espaciador_izq.pack(side="left")

        titulo_label = tk.Label(
            header,
            text="BASE DE DATOS",
            font=("Inter", 24, "bold"),
            fg="#000",
            bg=BG_COLOR
        )
        titulo_label.pack(side="left", expand=True)
        
        self.burger_btn = tk.Button(
            header, text="≡", font=("Inter", 18), bg=BG_COLOR, bd=0, cursor="hand2",
            command=self.mostrar_menu
        )
        self.burger_btn.pack(side="right", padx=(0, 20), pady=4)

        # --- Scroll general ---
        self.canvas = tk.Canvas(self, bg=BG_COLOR, highlightthickness=0, borderwidth=0)
        scrollbar = tk.Scrollbar(self, orient="vertical", command=self.canvas.yview)
        self.scrollable_frame = tk.Frame(self.canvas, bg=BG_COLOR)

        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self._ajustar_scrollregion()
        )
        self._canvas_window_id = self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.bind("<Configure>", lambda e: self.canvas.itemconfig(self._canvas_window_id, width=e.width))

        self.canvas.configure(yscrollcommand=scrollbar.set)
        self.canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # --- Rueda global: que siempre mueva el Canvas, desde donde estés ---
        self._wheel_target = self.canvas
        self.bind_all("<MouseWheel>", self._on_mousewheel)     # Windows / Mac
        self.bind_all("<Button-4>", self._on_mousewheel)       # Linux (scroll up)
        self.bind_all("<Button-5>", self._on_mousewheel)       # Linux (scroll down)

        # --- Cargar datos ---
        self.datos = self.obtener_datos()
        print("DEBUG total filas traídas:", len(self.datos))
        self.update_idletasks()
        
        # --- paginación (solo config) ---
        self.cards_per_page = 32  # <<<< clave: evitar superar ~32k px de alto por página
        self.current_page = 1

        # Render de la página actual
        self.mostrar_datos_agrupados()

        # --- DEBUG: cuántas tarjetas quedaron renderizadas ---
        try:
            total_cards = sum(
                1 for w in self.scrollable_frame.winfo_children()
                if isinstance(w, tk.Frame) and w is not None
                and w.cget("bg") == FRAME_COLOR
            )
            print("DEBUG tarjetas renderizadas:", total_cards)
        except Exception:
            pass

    # =================== DEBUG HELPERS ===================
    def _toggle_debug_overlay(self):
        self._debug_overlay_on = not self._debug_overlay_on
        self._debug_dump("F2 toggle")

    def _contenido_ymax(self):
        """y + alto del hijo más bajo (fondo real del contenido)."""
        ymax = 0
        for w in self.scrollable_frame.winfo_children():
            try:
                y = w.winfo_y()
                h = w.winfo_height()
                ymax = max(ymax, y + h)
            except Exception:
                pass
        return ymax

    def _encontrar_ultimo_card(self):
        """Encuentra el último frame de tarjeta (por estilo) y su fondo."""
        last_w = None
        last_bottom = 0
        for w in self.scrollable_frame.winfo_children():
            try:
                if isinstance(w, tk.Frame) and str(w.cget("bg")) == FRAME_COLOR:
                    y = w.winfo_y()
                    h = w.winfo_height()
                    if y + h >= last_bottom:
                        last_bottom = y + h
                        last_w = w
            except Exception:
                pass
        return last_w, last_bottom

    def _debug_dump(self, origen):
        try:
            self.update_idletasks()
            visible_h = self.canvas.winfo_height()
            req_h = self.scrollable_frame.winfo_reqheight()
            real_h = self.scrollable_frame.winfo_height()
            ymax = self._contenido_ymax()
            sr = self.canvas.cget("scrollregion") or "0 0 0 0"
            try:
                sr_parts = [int(float(x)) for x in sr.split()]
                sr_h = sr_parts[3] if len(sr_parts) == 4 else 0
            except Exception:
                sr_h = 0
            _, last_card_bottom = self._encontrar_ultimo_card()

            print(f"\n[DEBUG:{origen}] pag={self.current_page}  "
                  f"visible_h={visible_h}  req={req_h} real={real_h}  "
                  f"ymax={ymax}  last_card_bottom={last_card_bottom}  "
                  f"scrollregion_h={sr_h}")

            # Dibujar overlay (rojo=sr_h, verde=ymax, azul=última tarjeta)
            self.canvas.delete("dbg_lines")
            if self._debug_overlay_on:
                w = max(1, self.canvas.winfo_width())
                self.canvas.create_line(0, sr_h, w, sr_h, fill="red", width=2, tags="dbg_lines")
                self.canvas.create_line(0, ymax, w, ymax, fill="green", width=2, tags="dbg_lines")
                self.canvas.create_line(0, last_card_bottom, w, last_card_bottom, fill="blue", width=2, tags="dbg_lines")
        except Exception:
            pass
    # =====================================================

    # --- NUEVO: ajuste del scrollregion/medidas finales ---
    def _ajustar_scrollregion(self):
        """
        Fija el scrollregion usando el fondo real (ymax) del contenido.
        Agrega un margen chico solo si hay scroll real.
        """
        try:
            self.update_idletasks()
            try:
                self.canvas.itemconfig(self._canvas_window_id, width=self.canvas.winfo_width())
            except Exception:
                pass

            visible_h = self.canvas.winfo_height() or 1
            req_h = self.scrollable_frame.winfo_reqheight()
            real_h = self.scrollable_frame.winfo_height()
            ymax = self._contenido_ymax()

            content_h = max(ymax, real_h, req_h, 1)
            extra = 24 if content_h > visible_h else 0
            content_w = self.canvas.winfo_width() or 1
            self.canvas.configure(scrollregion=(0, 0, content_w, content_h + extra))

            self._debug_dump("_ajustar_scrollregion")
        except Exception:
            pass

    def _build_pagination(self, parent, total_pages):
        """Crea barra de paginación dentro de 'parent', centrada."""
        # limpiar contenido previo
        for w in parent.winfo_children():
            w.destroy()

        nav = tk.Frame(parent, bg=BG_COLOR)
        nav.pack(anchor="center")

        def add_label(txt, go_to=None, enabled=True, is_current=False):
            fg = "#1746A2" if enabled else "#999999"
            font = ("Inter", 12, "bold") if is_current else ("Inter", 12)
            lbl = tk.Label(nav, text=txt, bg=BG_COLOR, fg=fg,
                           font=font, cursor=("hand2" if enabled else "arrow"))
            lbl.pack(side="left", padx=6)
            if enabled and go_to is not None:
                lbl.bind("<Button-1>", lambda e, p=go_to: self._go_to_page(p))

        # total_pages mínimo 1
        total_pages = max(1, int(total_pages))

        # « primera, ‹ anterior
        add_label("«", go_to=1, enabled=(self.current_page > 1))
        add_label("‹", go_to=(self.current_page - 1), enabled=(self.current_page > 1))

        # números
        for p in range(1, total_pages + 1):
            add_label(str(p), go_to=p, enabled=(p != self.current_page), is_current=(p == self.current_page))

        # › siguiente, » última
        add_label("›", go_to=(self.current_page + 1), enabled=(self.current_page < total_pages))
        add_label("»", go_to=total_pages, enabled=(self.current_page < total_pages))

    def _go_to_page(self, page_num):
        """Cambiar página y re-renderizar."""
        self.current_page = max(1, int(page_num))
        self.mostrar_datos_agrupados()
        # ir al tope de la nueva página
        try:
            self.canvas.yview_moveto(0.0)
        except Exception:
            pass
        # recalcular scrollregion (inmediato + after_idle + pequeño delay)
        self._ajustar_scrollregion()
        self.after_idle(self._ajustar_scrollregion)
        self.after(60, self._ajustar_scrollregion)

    # Manejo universal de rueda (Windows/Mac/Linux)
    def _bind_wheel_to(self, widget):
        self._wheel_target = widget
        self.bind_all("<MouseWheel>", self._on_mousewheel)     # Windows / Mac
        self.bind_all("<Button-4>", self._on_mousewheel)       # Linux up
        self.bind_all("<Button-5>", self._on_mousewheel)       # Linux down)

    def _unbind_wheel(self):
        # Mantener la rueda siempre activa sobre el Canvas
        self._wheel_target = self.canvas  # no deshacer los bind_all

    def _on_mousewheel(self, event):
        c = self._wheel_target
        if not c:
            return

        # Normalizamos el "sentido" del scroll
        if hasattr(event, "num") and event.num in (4, 5):  # Linux
            steps = -1 if event.num == 4 else 1
        else:  # Windows / Mac
            # event.delta suele ser múltiplo de 120
            sign = -1 if event.delta > 0 else 1
            steps = sign * max(1, abs(event.delta) // 120)

        try:
            # En Canvas, scrolleo en PIXELES para poder llegar exactamente al fondo.
            if isinstance(c, tk.Canvas):
                c.yview_scroll(steps * 60, "pixels")   # ~60px por notch
            else:
                # En Text/ScrolledText, "units" (líneas) es lo correcto.
                c.yview_scroll(steps, "units")
        except Exception:
            pass

    def obtener_datos(self):
        try:
            page_size = 50  # si el server capea ~41 igual vamos pidiendo por tandas
            todos = []
            # usar keyset pagination con cursor por ID
            last_id = 0  # trae id > 0

            while True:
                resp = (
                    supabase
                    .table("compendio")
                    .select("*")
                    .order("id", desc=False)
                    .gt("id", last_id)
                    .limit(page_size)
                    .execute()
                )
                data = getattr(resp, "data", None) or []
                if not data:
                    break

                todos.extend(data)

                # debug del rango de IDs por tanda
                try:
                    print(f"DEBUG tanda ids: {data[0].get('id')} -> {data[-1].get('id')}")
                except Exception:
                    pass

                # avanzar cursor al último id del lote
                try:
                    last_id = data[-1].get("id")
                except Exception:
                    last_id = None

                if last_id is None:
                    break

                if len(data) < page_size:
                    break

            # orden final (como ya tenías)
            try:
                todos.sort(key=lambda r: ((r.get("tema") or ""), (r.get("subtema") or ""), r.get("id") or 0))
            except Exception:
                pass

            return todos

        except Exception as e:
            messagebox.showerror("Error", f"No se pudo obtener la información: {str(e)}")
            return []

    # ======== RENDER PAGINADO ========
    def mostrar_datos_agrupados(self):
        self.limpiar_vista()

        # --- slicing por página ---
        total_records = len(self.datos)
        total_pages = max(1, (total_records + self.cards_per_page - 1) // self.cards_per_page)
        self.current_page = max(1, min(self.current_page, total_pages))
        start = (self.current_page - 1) * self.cards_per_page
        end = start + self.cards_per_page
        subset = self.datos[start:end]

        # --- barra de paginación ARRIBA (debajo del título, antes de la 1ra tarjeta) ---
        self.pagination_top = tk.Frame(self.scrollable_frame, bg=BG_COLOR)
        self.pagination_top.pack(fill="x", pady=(6, 8))
        self._build_pagination(self.pagination_top, total_pages)

        # --- agrupar SOLO lo de esta página ---
        agrupado = {}
        for reg in subset:
            tema = reg.get("tema") or "SIN TEMA"
            subtema = reg.get("subtema") or "SIN SUBTEMA"
            agrupado.setdefault(tema, {})
            agrupado[tema].setdefault(subtema, [])
            agrupado[tema][subtema].append(reg)

        # contador global visible en cada tarjeta
        idx_global = start + 1

        for tema in sorted(agrupado):
            tk.Label(
                self.scrollable_frame,
                text=tema,
                font=TITULO_TEMA,
                bg=BG_COLOR,
                fg="#000"
            ).pack(anchor="w", pady=(28, 4), padx=30)

            for subtema in sorted(agrupado[tema]):
                if subtema and subtema != "SIN SUBTEMA":
                    tk.Label(
                        self.scrollable_frame,
                        text=f"» {subtema}",
                        font=TITULO_SUBTEMA,
                        bg=BG_COLOR,
                        fg="#000"
                    ).pack(anchor="w", padx=52)

                for reg in agrupado[tema][subtema]:
                    self.visualizar_registro(reg, idx_global)  # << numeración
                    idx_global += 1

        # --- barra de paginación ABAJO (debajo de la última tarjeta) ---
        # contenedor de ancho completo para que el canvas lo mida bien
        self.pagination_bottom_outer = tk.Frame(
            self.scrollable_frame,
            bg=BG_COLOR,
            bd=0,
            highlightthickness=0
        )
        self.pagination_bottom_outer.pack(side="top", fill="x")

        # barra centrada dentro del contenedor
        self.pagination_bottom = tk.Frame(
            self.pagination_bottom_outer,
            bg=BG_COLOR,
            bd=0,
            highlightthickness=0
        )
        self.pagination_bottom.pack(side="top", pady=(12, 16), anchor="center")

        # construir la barra de paginación
        self._build_pagination(self.pagination_bottom, total_pages)

        # --- ESPACIADOR FINAL (colchón grande y dinámico) ---
        if hasattr(self, "_end_spacer") and self._end_spacer.winfo_exists():
            self._end_spacer.destroy()

        self.update_idletasks()
        try:
            visible_h = self.canvas.winfo_height()
        except Exception:
            visible_h = 600  # fallback

        spacer_h = max(visible_h // 2, 160)

        self._end_spacer = tk.Frame(
            self.scrollable_frame,
            height=spacer_h,
            bg=BG_COLOR,
            bd=0,
            highlightthickness=0
        )
        self._end_spacer.pack(side="top", fill="x")

        # Recalcular scrollregion cuando TODO ya está colocado y estabilizado
        self._ajustar_scrollregion()
        self.after_idle(self._ajustar_scrollregion)
        self.after(60, self._ajustar_scrollregion)

        # DEBUG: volcar métricas al final del render
        self._debug_dump("render_end")

    def visualizar_registro(self, reg, numero):
        try:
            card = tk.Frame(
                self.scrollable_frame,
                bg=FRAME_COLOR,
                bd=1,                 # borde físico parejo (sin highlight)
                relief="solid",
                highlightthickness=0  # evita la “rayita” por highlight recortado
            )
            # clave: no expand=True
            card.pack(fill="x", padx=38, pady=18, anchor="center", expand=False)
            
            # mantener scroll de la vista
            card.bind("<Enter>", lambda e: self._bind_wheel_to(self.canvas))
            card.bind("<Leave>", lambda e: self._bind_wheel_to(self.canvas))

            card.grid_columnconfigure(0, weight=9)
            card.grid_columnconfigure(1, weight=1)

            info_frame = tk.Frame(card, bg=FRAME_COLOR)
            info_frame.grid(row=0, column=0, sticky="nw", padx=(0, 18), pady=(4, 4))

            # SANEAMOS: todo a string seguro
            def s(v):
                if v is None:
                    return "-"
                try:
                    v = str(v)
                except Exception:
                    v = "-"
                return v if v.strip() != "" else "-"

            sections = [
                ("AUTOS CARATULADOS", s(reg.get("autos", ""))),
                ("JURISDICCIÓN/INSTANCIA", s(reg.get("jurisdiccion", ""))),
                ("FECHA DE SENTENCIA", s(reg.get("fecha", ""))),
                ("RESULTADO", s(reg.get("resultado", ""))),
                ("VOCES", s(reg.get("voces", ""))),
                ("LINK FALLO", reg.get("link_fallo", "")),  # tratamos aparte
                ("JURISPRUDENCIA", s(reg.get("jurisprudencia", ""))),
            ]

            for label, valor in sections:
                block = tk.Frame(info_frame, bg=FRAME_COLOR)
                block.pack(anchor="w", pady=(1,4), fill="x")
                tk.Label(
                    block,
                    text=label + ":",
                    font=LABEL_FONT,
                    fg="#1746A2",
                    bg=FRAME_COLOR
                ).pack(anchor="w")

                if label == "LINK FALLO":
                    # solo tratamos como link si es string no-vacío
                    if isinstance(valor, str) and valor.strip() not in ["", "-"]:
                        enlace_texto = valor.strip()
                        enlace = tk.Label(
                            block,
                            text=enlace_texto,
                            font=(INFO_FONT[0], INFO_FONT[1], "underline"),
                            fg="#1662bb",
                            bg=FRAME_COLOR,
                            cursor="hand2",
                            wraplength=1220,
                            justify="left"
                        )
                        enlace.pack(anchor="w")
                        enlace.bind(
                            "<Button-1>",
                            lambda e, url=enlace_texto: webbrowser.open_new(url if isinstance(url, str) and url.startswith("http") else f"https://{url}")
                        )
                    else:
                        tk.Label(
                            block,
                            text=s(valor),
                            font=INFO_FONT,
                            fg="#222",
                            bg=FRAME_COLOR,
                            wraplength=1220,
                            justify="left"
                        ).pack(anchor="w")
                elif label == "JURISPRUDENCIA":
                    st = ScrolledText(
                        block,
                        font=INFO_FONT,
                        wrap="word",
                        height=JURIS_ALTURA_LINEAS,
                        bd=1,
                        relief="solid"
                    )
                    st.insert("1.0", s(valor))
                    st.configure(state="disabled")
                    st.pack(fill="x", padx=0, pady=(0, 0))
                    st.bind("<Enter>", lambda e, c=st: self._bind_wheel_to(c))
                    st.bind("<Leave>", lambda e: self._bind_wheel_to(self.canvas))
                else:
                    tk.Label(
                        block,
                        text=s(valor),
                        font=INFO_FONT,
                        fg="#222",
                        bg=FRAME_COLOR,
                        wraplength=1220,
                        justify="left"
                    ).pack(anchor="w")

            # Botones
            btns_frame = tk.Frame(card, bg=FRAME_COLOR)
            btns_frame.grid(row=0, column=1, sticky="ne", padx=(0, 8), pady=10)
            btn_style = dict(
                font=BTN_ICONO,
                bg=BTN_COLOR,
                fg=BTN_TEXT_COLOR,
                activebackground=BTN_HOVER_BG,
                activeforeground=BTN_HOVER_TEXT,
                relief="flat",
                borderwidth=0,
                highlightthickness=2,
                highlightbackground=BTN_BORDER_COLOR,
                cursor="hand2",
                width=3,
                height=1,
            )
            btn_copiar = tk.Button(
                btns_frame,
                text=ICON_COPIAR,
                command=lambda t=self.armar_texto_copiar(reg): self.copiar_a_clipboard(t),
                **btn_style
            )
            btn_copiar.pack(pady=7)
            btn_descargar = tk.Button(
                btns_frame,
                text=ICON_DESCARGAR,
                command=lambda d=reg: self.descargar_docx(d),
                **btn_style
            )
            btn_descargar.pack(pady=7)
            btn_editar = tk.Button(
                btns_frame,
                text=ICON_EDITAR,
                command=lambda r=reg: self.editar_registro(r),
                **btn_style
            )
            btn_editar.pack(pady=7)

            btn_borrar = tk.Button(
                btns_frame,
                image=self.icono_tacho_azul,
                command=lambda r=reg: self.borrar_registro(r),
                bg=BTN_COLOR,
                activebackground=BTN_HOVER_BG,
                relief="flat",
                borderwidth=0,
                highlightthickness=2,
                highlightbackground=BTN_BORDER_COLOR,
                cursor="hand2",
                width=34,
                height=34
            )
            btn_borrar.pack(pady=7)

            for btn in [btn_copiar, btn_descargar, btn_editar, btn_borrar]:
                btn.bind("<Enter>", lambda e, b=btn: b.config(bg=BTN_HOVER_BG))
                btn.bind("<Leave>", lambda e, b=btn: b.config(bg=BTN_COLOR))

            # --- NÚMERO DE TARJETA (arriba derecha) ---
            card.update_idletasks()  # asegura tamaño final del card
            badge = tk.Label(
                card,
                text=str(numero),
                bg=FRAME_COLOR,
                fg="#1746A2",
                font=("Inter", 10, "bold")
            )
            badge.place(relx=1.0, x=-12, y=6, anchor="ne")
            badge.lift()  # trae al frente por si algo lo tapa

        except Exception as ex:
            # No frenamos el render del resto
            print("ERROR al renderizar id=", reg.get("id"), " -> ", repr(ex))
            # seguimos (no re-lanzamos)

    def armar_texto_copiar(self, reg):
        def s(v):
            if v is None:
                return "-"
            try:
                v = str(v)
            except Exception:
                v = "-"
            return v

        campos = [
            ("AUTOS CARATULADOS", s(reg.get("autos", ""))),
            ("JURISDICCIÓN/INSTANCIA", s(reg.get("jurisdiccion", ""))),
            ("FECHA DE SENTENCIA", s(reg.get("fecha", ""))),
            ("RESULTADO", s(reg.get("resultado", ""))),
            ("VOCES", s(reg.get("voces", ""))),
            ("LINK FALLO", s(reg.get("link_fallo", ""))),
            ("JURISPRUDENCIA", s(reg.get("jurisprudencia", ""))),
        ]
        return "\n".join(f"{k}: {v}" for k, v in campos if v is not None)

    def copiar_a_clipboard(self, texto):
        self.clipboard_clear()
        self.clipboard_append(texto)
        # Silencioso, sin messagebox
    
    def mostrar_menu(self):
        x = self.burger_btn.winfo_rootx() - self.menu_popup.winfo_reqwidth()
        y = self.burger_btn.winfo_rooty() + self.burger_btn.winfo_height()
        x = max(10, x)
        self.menu_popup.tk_popup(x, y)
        self.menu_popup.grab_release()

    def descargar_docx(self, datos):
        doc = Document()
        doc.add_heading(f'Tema: {datos.get("tema","-")}  -  Subtema: {datos.get("subtema","-")}', level=1)
        doc.add_paragraph(f'AUTOS CARATULADOS:\n{datos.get("autos", "")}')
        doc.add_paragraph(f'JURISDICCIÓN/INSTANCIA:\n{datos.get("jurisdiccion", "")}')
        doc.add_paragraph(f'FECHA DE SENTENCIA: {datos.get("fecha", "-")}   RESULTADO: {datos.get("resultado", "-")}')
        doc.add_paragraph(f'VOCES:\n{datos.get("voces", "-")}')
        doc.add_paragraph(f'LINK FALLO:\n{datos.get("link_fallo", "-")}')
        doc.add_paragraph("\nJURISPRUDENCIA:")
        doc.add_paragraph(datos.get("jurisprudencia", "-"))

        archivo_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Documentos Word", "*.docx")],
            title="Guardar documento"  # opcional, solo cambia el título del diálogo
            # sin initialfile -> no sugiere ningún nombre
        )
        if archivo_path:
            doc.save(archivo_path)

    def limpiar_vista(self):
        for widget in self.scrollable_frame.winfo_children():
            widget.destroy()
        # asegurar que el canvas se entere del “vacío/lleno”
        self._ajustar_scrollregion()

    def editar_registro(self, reg):
        password = simpledialog.askstring("Editar registro", "Ingrese la contraseña para editar:", show="*")
        if password != ADMIN_PASSWORD:
            messagebox.showerror("Error", "Contraseña incorrecta.")
            return

        def guardar_edicion(datos_editados):
            try:
                supabase.table("compendio").update(datos_editados).eq("id", reg["id"]).execute()
                messagebox.showinfo("Éxito", "Registro actualizado correctamente.")
                self.datos = self.obtener_datos()
                self.limpiar_vista()
                self.mostrar_datos_agrupados()
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo actualizar: {str(e)}")

        # --- PASA el callback para restaurar menú principal ---
        menu_callback = self.master.deiconify if hasattr(self.master, "deiconify") else None
        EditaRegistro(self, reg, guardar_edicion, volver_menu_callback=menu_callback)

    # --- NUEVO: borrar registro individual con contraseña ---
    def borrar_registro(self, reg):
        password = simpledialog.askstring("Borrar registro", "Ingrese la contraseña para borrar:", show="*")
        if password != ADMIN_PASSWORD:
            messagebox.showerror("Error", "Contraseña incorrecta.")
            return
        if not messagebox.askyesno("Confirmar borrado", "¿Seguro que desea borrar este registro? Esta acción no se puede deshacer."):
            return
        try:
            supabase.table("compendio").delete().eq("id", reg["id"]).execute()
            messagebox.showinfo("Éxito", "Registro borrado correctamente.")
            self.datos = self.obtener_datos()
            self.limpiar_vista()
            self.mostrar_datos_agrupados()
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo borrar: {str(e)}")

    def busqueda_avanzada(self):
        self.iconify()  # Minimiza la ventana de información
        buscador = BusquedaAvanzada(self, volver_callback=self.deiconify)
        buscador.protocol("WM_DELETE_WINDOW", lambda: [buscador.destroy(), self.deiconify()])

    def volver_menu(self):
        # Primero restaurar el menú principal
        if self.master is not None and hasattr(self.master, 'deiconify'):
            self.master.deiconify()
        self.destroy()

    def quit_app(self):
        self.destroy()

# Para probar directo:
if __name__ == "__main__":
    root = tk.Tk()
    root.withdraw()
    VisualizadorBase(root)
    root.mainloop()
